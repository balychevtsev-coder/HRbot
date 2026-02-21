import docx

def extract_resume_data_from_docx(docx_bytes):
    from docx import Document
    from io import BytesIO
    doc = Document(BytesIO(docx_bytes))
    full_text = []
    
    # 1. Сначала читаем параграфы (заголовки и т.д.)
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # 2. Читаем таблицы (критично для требований!)
    for table in doc.tables:
        for row in table.rows:
            # Очищаем ячейки от лишних пробелов
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) > 1:
                full_text.append(f"{cells[0]}: {cells[1]}")
            elif len(cells) == 1:
                full_text.append(cells[0])
                
    return "\n".join(full_text)